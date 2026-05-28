from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "MediReservas-Informe-Tecnico.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(25, 32, 43)
MUTED = RGBColor(88, 101, 117)
TABLE_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table):
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MediReservas")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma de Reservas para Consultas Medicas")
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Arquitectura de Microservicios | MVP 1.0")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    data = [
        ("Elemento", "Detalle"),
        ("Proyecto", "MediReservas"),
        ("Arquitectura", "Microservicios"),
        ("Stack", "Node.js, Express, PostgreSQL, NGINX, Docker Compose"),
        ("Ejecucion", "docker compose up --build"),
    ]
    for row, values in zip(table.rows, data):
        for cell, value in zip(row.cells, values):
            cell.text = value
        set_cell_shading(row.cells[0], TABLE_FILL)
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, TABLE_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    doc.add_paragraph()


def add_architecture_diagram(doc):
    doc.add_heading("Diagrama de arquitectura general", level=2)
    add_simple_table(
        doc,
        ["Entrada", "Gateway", "Microservicios", "Persistencia"],
        [
            (
                "Paciente / Medico",
                "NGINX API Gateway\nPuertos 8081 y 8082",
                "Auth Service\nAgenda Service\nNotifications Service\nRecords Service",
                "PostgreSQL\nTablas por dominio",
            )
        ],
    )


def add_extra_diagrams(doc):
    doc.add_heading("Diagrama de casos de uso", level=2)
    add_simple_table(
        doc,
        ["Actor", "Casos de uso"],
        [
            ("Paciente", "Registrarse, iniciar sesion, consultar medicos disponibles, agendar, modificar, reprogramar y cancelar citas, consultar notificaciones y resultados."),
            ("Medico", "Iniciar sesion, configurar perfil medico, publicar disponibilidad, consultar agenda, listar/agregar pacientes y registrar resultados."),
        ],
    )
    doc.add_heading("Diagrama de secuencia: agendar cita", level=2)
    add_numbered(
        doc,
        [
            "Paciente envia POST /api/agenda/appointments al API Gateway.",
            "Agenda Service valida el token con Auth Service.",
            "Agenda Service bloquea el horario disponible en PostgreSQL.",
            "Agenda Service crea la cita y marca el horario como reservado.",
            "Agenda Service solicita a Notifications Service crear una notificacion.",
            "El paciente recibe la confirmacion de la cita.",
        ],
    )


def add_screenshot(doc, title, image_name):
    path = ROOT / "docs" / image_name
    if path.exists():
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.3))


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("Descripcion del problema", level=1)
    doc.add_paragraph(
        "Las clinicas pequenas suelen coordinar citas por telefono o mensajeria, lo que provoca doble reserva de horarios, baja trazabilidad, dificultad para administrar disponibilidad medica y poca visibilidad para el paciente. MediReservas centraliza el proceso con servicios desacoplados."
    )

    doc.add_heading("Alcance del MVP", level=1)
    doc.add_paragraph(
        "El MVP permite registrar usuarios, iniciar sesion, configurar perfil medico, listar medicos disponibles, publicar disponibilidad, agendar, modificar, reprogramar y cancelar citas, emitir notificaciones internas, agregar pacientes y registrar/consultar resultados basicos de consulta."
    )

    doc.add_heading("Requerimientos funcionales", level=1)
    add_bullets(
        doc,
        [
            "Registro e inicio de sesion de usuarios.",
            "Roles de paciente, medico y administrador.",
            "Configuracion de perfil medico.",
            "Publicacion de bloques de disponibilidad.",
            "Consulta de medicos disponibles, contador de horarios y proximo horario.",
            "Agendamiento, modificacion, reprogramacion y cancelacion de citas.",
            "Listado y registro de pacientes desde el portal medico.",
            "Generacion de notificaciones internas.",
            "Registro y consulta de resultados medicos basicos.",
        ],
    )

    doc.add_heading("Requerimientos no funcionales", level=1)
    add_bullets(
        doc,
        [
            "Modularidad por dominio y contratos HTTP.",
            "Ejecucion local reproducible con Docker Compose.",
            "Seguridad mediante token firmado.",
            "Trazabilidad con estados y fechas de creacion.",
            "Consistencia de agenda mediante transaccion SQL y bloqueo de horario.",
            "Mantenibilidad por separacion de responsabilidades.",
        ],
    )

    doc.add_heading("Arquitectura y diagramas", level=1)
    doc.add_paragraph(
        "La solucion usa NGINX como API Gateway y punto de entrada unico. El gateway enruta llamadas hacia Auth, Agenda, Notifications y Records. Los servicios se comunican por HTTP y persisten datos en PostgreSQL para simplificar la ejecucion local del MVP."
    )
    add_architecture_diagram(doc)
    add_extra_diagrams(doc)

    doc.add_heading("Stack tecnologico", level=1)
    add_simple_table(
        doc,
        ["Capa", "Tecnologia", "Justificacion"],
        [
            ("Gateway", "NGINX", "Entrada unificada y reverse proxy."),
            ("Servicios", "Node.js + Express", "Ligero y adecuado para APIs REST."),
            ("Base de datos", "PostgreSQL 16", "Transacciones ACID para evitar doble reserva."),
            ("Frontend", "HTML, CSS y JavaScript", "Interfaz simple para demostrar el MVP."),
            ("Infraestructura", "Docker Compose", "Ejecucion multi-contenedor local."),
        ],
    )

    doc.add_heading("Modelo de datos", level=1)
    add_simple_table(
        doc,
        ["Entidad", "Responsabilidad"],
        [
            ("auth_users", "Usuarios, credenciales y roles."),
            ("agenda_doctors", "Perfil profesional del medico."),
            ("agenda_availability_slots", "Bloques de disponibilidad."),
            ("agenda_appointments", "Citas entre paciente y medico."),
            ("notification_messages", "Notificaciones internas."),
            ("record_results", "Resultados o notas de consulta."),
        ],
    )

    doc.add_heading("Decisiones tecnicas y trade-offs", level=1)
    add_bullets(
        doc,
        [
            "NGINX se usa como gateway para mantener un punto de entrada claro.",
            "La comunicacion entre servicios es HTTP sincrona por simplicidad del MVP.",
            "Se usa una sola instancia PostgreSQL local con tablas por dominio; en produccion podria separarse por servicio.",
            "Auth Service centraliza validacion de tokens para no duplicar seguridad.",
            "Agenda Service usa transacciones y bloqueo de fila para prevenir doble reserva.",
        ],
    )

    doc.add_heading("Evidencia del sistema", level=1)
    add_bullets(
        doc,
        [
            "Portal paciente disponible en http://localhost:8081.",
            "Portal medico disponible en http://localhost:8082.",
            "Smoke test automatizado disponible en scripts/smoke-test.ps1.",
            "Contenedores verificados: postgres, auth, agenda, notifications, records y gateway.",
        ],
    )
    add_screenshot(doc, "Captura: portal paciente", "evidencia-paciente.png")
    add_screenshot(doc, "Captura: portal medico", "evidencia-medico.png")

    doc.add_heading("Instrucciones de ejecucion", level=1)
    add_numbered(
        doc,
        [
            'Abrir PowerShell y ejecutar: cd "C:\\Users\\USUARIO\\Documents\\fernando\\fernando proyecto analisis".',
            "Ejecutar: docker compose up --build.",
            "Abrir portal paciente en http://localhost:8081.",
            "Abrir portal medico en http://localhost:8082.",
            "Ejecutar prueba automatizada: .\\scripts\\smoke-test.ps1.",
        ],
    )

    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(
        "MediReservas evidencia separacion de responsabilidades, contratos de API, comunicacion entre servicios, despliegue local con contenedores y una solucion ejecutable coherente con microservicios."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

